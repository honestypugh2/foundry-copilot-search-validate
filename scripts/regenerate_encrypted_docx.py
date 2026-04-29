"""
Regenerate CDFV2-Encrypted .docx files with synthetic HR policy content.

Scans data/knowledge_base_lab/ for .docx files that are CDFV2 Encrypted
(not valid OOXML), generates plausible HR policy content based on the
filename, and overwrites the encrypted file with a proper .docx.

Usage:
    uv run python src/scripts/regenerate_encrypted_docx.py
    uv run python src/scripts/regenerate_encrypted_docx.py --dry-run
"""

import argparse
import logging
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

DATA_DIR = Path(__file__).resolve().parents[2] / "data" / "knowledge_base_lab"

# ---------------------------------------------------------------------------
# Content templates keyed by topic patterns found in filenames
# ---------------------------------------------------------------------------

POLICY_CONTENT = {
    "code of ethics": {
        "purpose": "To establish standards of ethical conduct for all employees and ensure compliance with applicable laws and regulations.",
        "scope": "This policy applies to all full-time, part-time, and temporary employees of the organization.",
        "sections": [
            ("Standards of Conduct", "All employees are expected to conduct themselves with integrity, honesty, and professionalism in all business dealings. Employees must avoid conflicts of interest and disclose any potential conflicts to their supervisor or the Ethics Committee."),
            ("Conflicts of Interest", "Employees must not engage in activities that conflict with the interests of the organization. This includes financial interests in competing businesses, accepting gifts from vendors exceeding $25 in value, and using company resources for personal gain."),
            ("Reporting Violations", "Employees who become aware of violations of this Code of Ethics are encouraged to report them through the Ethics Hotline or directly to the Ethics Committee. Reports may be made anonymously. The organization prohibits retaliation against employees who report violations in good faith."),
            ("Compliance and Enforcement", "Violations of this Code of Ethics may result in disciplinary action, up to and including termination of employment. The Ethics Committee is responsible for investigating reported violations and recommending appropriate corrective action."),
        ],
    },
    "pre-employment medical": {
        "purpose": "To outline the requirements and procedures for pre-employment medical examinations as a condition of employment.",
        "scope": "This policy applies to all candidates who have received a conditional offer of employment.",
        "sections": [
            ("Examination Requirements", "All candidates receiving a conditional job offer must complete a pre-employment medical examination conducted by an authorized healthcare provider. The examination must be completed within 30 days of the conditional offer."),
            ("Fitness for Duty", "The medical examination will assess the candidate's ability to perform the essential functions of the position with or without reasonable accommodation. Results are confidential and maintained in separate medical files."),
            ("Reasonable Accommodation", "If a medical examination reveals a condition that may affect job performance, the organization will engage in an interactive process to determine if reasonable accommodation can be provided in accordance with the Americans with Disabilities Act (ADA)."),
            ("Confidentiality", "All medical information obtained through pre-employment examinations is treated as confidential and maintained in compliance with HIPAA and applicable state privacy laws."),
        ],
    },
    "rehiring.*retirees": {
        "purpose": "To establish guidelines for the rehiring of retired employees without the requirement of advertising the position.",
        "scope": "This policy applies to former employees who retired in good standing and are being considered for re-employment.",
        "sections": [
            ("Eligibility", "Retired employees who separated from the organization in good standing and who possess current skills relevant to available positions may be considered for re-employment without advertising the position externally."),
            ("Approval Process", "Department heads must submit a written request to Human Resources documenting the business need and the retiree's qualifications. Final approval rests with the Chief Human Resources Officer."),
            ("Terms of Re-employment", "Re-employed retirees may be hired as regular full-time, part-time, or temporary employees. Benefits eligibility is determined based on the terms of re-employment and applicable benefit plan provisions."),
            ("Impact on Retirement Benefits", "Re-employed retirees should consult with the Benefits Office regarding any impact on their retirement benefits, including pension payments and retiree health insurance coverage."),
        ],
    },
    "probationary period": {
        "purpose": "To define the probationary period for newly hired and promoted employees and establish performance expectations during this period.",
        "scope": "This policy applies to all newly hired employees and employees who have been promoted to a new position.",
        "sections": [
            ("Duration", "The standard probationary period is six (6) months from the date of hire or promotion. The probationary period may be extended for up to an additional three (3) months with approval from the Department Head and Human Resources."),
            ("Performance Expectations", "During the probationary period, employees are expected to demonstrate competency in the essential functions of their position, maintain satisfactory attendance, and comply with all organizational policies and procedures."),
            ("Evaluation", "Supervisors must conduct a formal performance evaluation at the midpoint and conclusion of the probationary period using the standard evaluation form. Evaluations must be submitted to Human Resources within five (5) business days of completion."),
            ("Completion and Separation", "Successful completion of the probationary period results in regular employment status. Employees who fail to meet performance standards during the probationary period may be separated from employment without the progressive discipline process."),
        ],
    },
    "holiday pay": {
        "purpose": "To establish the organization's policy on holiday pay and designate official paid holidays.",
        "scope": "This policy applies to all regular full-time employees. Part-time employees receive holiday pay on a pro-rated basis.",
        "sections": [
            ("Designated Holidays", "The organization observes the following paid holidays: New Year's Day, Martin Luther King Jr. Day, Presidents' Day, Memorial Day, Independence Day, Labor Day, Columbus Day, Veterans Day, Thanksgiving Day, Day after Thanksgiving, Christmas Eve, and Christmas Day."),
            ("Holiday Pay Rate", "Eligible employees receive their regular rate of pay for each designated holiday. Employees required to work on a designated holiday receive holiday pay plus time-and-a-half for hours worked."),
            ("Holiday Scheduling", "When a designated holiday falls on a Saturday, the preceding Friday is observed. When a holiday falls on a Sunday, the following Monday is observed."),
            ("Eligibility Requirements", "To receive holiday pay, employees must work their regularly scheduled shift on the workday immediately before and after the holiday, unless on approved leave."),
        ],
    },
    "career path.*hr generalist": {
        "purpose": "To define the career progression path for Human Resources Generalist positions within the organization.",
        "scope": "This policy applies to all employees in the HR Generalist career track.",
        "sections": [
            ("Career Levels", "The HR Generalist career path includes the following levels: HR Assistant, HR Generalist I, HR Generalist II, Senior HR Generalist, and HR Manager. Each level has defined competency requirements and salary ranges."),
            ("Competency Requirements", "Progression through the career path requires demonstrated competency in core HR functions including recruitment, employee relations, benefits administration, compliance, and HRIS management. Specific competency matrices are maintained by the HR Department."),
            ("Education and Certification", "HR Generalist I requires a bachelor's degree in HR or related field. Senior HR Generalist requires PHR or SHRM-CP certification. HR Manager requires SPHR or SHRM-SCP certification."),
            ("Advancement Process", "Advancement to the next career level requires a minimum of two years in the current position, satisfactory performance evaluations, completion of required competencies, and approval from the CHRO."),
        ],
    },
    "career path.*data management": {
        "purpose": "To define the career progression path for Data Management positions within the organization.",
        "scope": "This policy applies to all employees in the Data Management career track.",
        "sections": [
            ("Career Levels", "The Data Management career path includes: Data Entry Clerk, Data Analyst I, Data Analyst II, Senior Data Analyst, and Data Management Supervisor. Each level has defined technical competency requirements."),
            ("Technical Competencies", "Progression requires demonstrated proficiency in database management, data quality assurance, reporting tools, and data governance practices. Specific technical certifications may be required at senior levels."),
            ("Training Requirements", "Each career level has mandatory training requirements that must be completed within 12 months of promotion. Training includes both technical skills and organizational knowledge."),
            ("Performance Standards", "Data accuracy rates, turnaround times, and compliance with data governance standards are key performance metrics at all career levels."),
        ],
    },
    "paid time off.*pto(?!.*part)": {
        "purpose": "To establish the Paid Time Off (PTO) policy for eligible full-time employees.",
        "scope": "This policy applies to all regular full-time employees who have completed their probationary period.",
        "sections": [
            ("PTO Accrual", "Full-time employees accrue PTO based on years of service: 0-2 years: 15 days per year, 3-5 years: 20 days per year, 6-10 years: 25 days per year, 11+ years: 30 days per year. PTO accrues on a per-pay-period basis."),
            ("PTO Usage", "PTO may be used for vacation, personal time, illness, or any other purpose. Employees must request PTO through the designated time-off system. Requests for three or more consecutive days require supervisor approval at least two weeks in advance."),
            ("Carryover and Maximum Balance", "Employees may carry over up to 40 hours of unused PTO to the following calendar year. Maximum PTO balance at any time is 240 hours. Accrual stops when the maximum balance is reached."),
            ("PTO at Separation", "Employees who separate from employment in good standing will be paid for unused accrued PTO up to a maximum of 80 hours at their current rate of pay."),
        ],
    },
    "paid time off.*part.time": {
        "purpose": "To establish the Paid Time Off (PTO) policy for eligible part-time employees.",
        "scope": "This policy applies to all regular part-time employees who work a minimum of 20 hours per week and have completed their probationary period.",
        "sections": [
            ("PTO Accrual for Part-Time Employees", "Part-time employees accrue PTO on a pro-rated basis calculated from their regularly scheduled hours. Accrual rates are: 0-2 years: 0.0385 hours per hour worked, 3-5 years: 0.0577 hours per hour worked, 6+ years: 0.0769 hours per hour worked."),
            ("Usage and Scheduling", "Part-time employees must request PTO in accordance with their regular schedule. PTO is deducted in hourly increments. Requests must be submitted through the time-off system."),
            ("Carryover", "Part-time employees may carry over up to 20 hours of unused PTO per calendar year. Accrual stops when maximum balance is reached."),
            ("Eligibility Changes", "Part-time employees who transition to full-time status will have their PTO balance converted and begin accruing under the full-time PTO policy."),
        ],
    },
    "short.term disability": {
        "purpose": "To provide income protection for eligible employees who are unable to work due to a qualifying short-term disability.",
        "scope": "This policy applies to all regular full-time employees who have completed 90 days of continuous employment.",
        "sections": [
            ("Benefit Amount", "Eligible employees receive 60% of their base weekly salary during an approved short-term disability period, subject to a maximum weekly benefit. Benefits begin after a seven (7) calendar day elimination period."),
            ("Duration of Benefits", "Short-term disability benefits are payable for up to 26 weeks per disability period. Benefits cease when the employee is able to return to work, reaches maximum benefit duration, or qualifies for long-term disability."),
            ("Qualifying Conditions", "Short-term disability covers non-work-related injuries and illnesses that prevent the employee from performing the essential functions of their position. Work-related injuries are covered under Workers' Compensation."),
            ("Claim Process", "Employees must notify their supervisor and Human Resources within three (3) business days of the onset of disability. A completed claim form and supporting medical documentation must be submitted to the disability administrator within 15 days."),
        ],
    },
    "uniform dress code": {
        "purpose": "To establish standards for employee uniforms and maintain a professional and consistent appearance across the organization.",
        "scope": "This policy applies to all employees who are required to wear uniforms as a condition of their position.",
        "sections": [
            ("Uniform Requirements", "Employees in designated positions must wear the approved uniform during all working hours. Uniforms consist of organization-issued shirts, pants or skirts, and approved footwear as specified for each job classification."),
            ("Uniform Issuance", "New employees receive an initial uniform allotment upon hire. Replacement uniforms are provided annually or as needed due to wear. Employees must return all uniform items upon separation from employment."),
            ("Maintenance and Appearance", "Employees are responsible for maintaining their uniforms in clean, pressed, and good repair condition. Uniforms that are damaged, stained, or excessively worn must be replaced through the uniform coordinator."),
            ("Modifications and Accessories", "Uniforms may not be altered or modified without approval. Name badges must be worn visibly at all times. Only organization-approved accessories may be worn with the uniform."),
        ],
    },
    "non.uniform dress code": {
        "purpose": "To establish professional dress and appearance standards for employees who are not required to wear uniforms.",
        "scope": "This policy applies to all employees who are not assigned a uniform.",
        "sections": [
            ("Business Professional Attire", "Employees are expected to dress in business professional attire appropriate to their work environment and job responsibilities. Business professional includes suits, dress shirts, blouses, slacks, skirts, and closed-toe dress shoes."),
            ("Business Casual", "On designated business casual days or in approved departments, employees may wear collared shirts, khakis, dress jeans (no rips or tears), and clean casual shoes. T-shirts, athletic wear, and flip-flops are not acceptable."),
            ("Grooming Standards", "Employees must maintain a clean and well-groomed appearance. Hair must be neat and professional. Visible tattoos and piercings beyond standard ear piercings must comply with departmental guidelines."),
            ("Exceptions and Accommodations", "Employees may request exceptions for religious, medical, or cultural reasons. Requests should be directed to Human Resources for review and accommodation."),
        ],
    },
    "acceptable use": {
        "purpose": "To define acceptable use of information technology resources provided by the organization.",
        "scope": "This policy applies to all employees, contractors, and authorized users of the organization's IT systems and resources.",
        "sections": [
            ("Authorized Use", "IT resources are provided for business purposes. Limited personal use is permitted provided it does not interfere with job performance, consume excessive resources, or violate any organizational policy. All users must acknowledge this policy annually."),
            ("Prohibited Activities", "Prohibited activities include: accessing or distributing inappropriate content, installing unauthorized software, sharing login credentials, circumventing security controls, using IT resources for personal financial gain, and sending mass unsolicited emails."),
            ("Monitoring and Privacy", "The organization reserves the right to monitor, access, and review all data stored on or transmitted through its IT systems. Users should have no expectation of privacy when using organizational IT resources."),
            ("Security Responsibilities", "Users must protect their login credentials, lock workstations when unattended, report security incidents immediately, and comply with all password and data protection requirements."),
            ("Enforcement", "Violations of this policy may result in disciplinary action up to and including termination, as well as potential civil or criminal liability."),
        ],
    },
    "information security": {
        "purpose": "To establish the organization's information security program and protect the confidentiality, integrity, and availability of information assets.",
        "scope": "This policy applies to all employees, contractors, and third parties who access or handle the organization's information assets.",
        "sections": [
            ("Information Classification", "Information is classified into four categories: Public, Internal, Confidential, and Restricted. Each classification level has defined handling, storage, and transmission requirements."),
            ("Access Control", "Access to information systems is granted on a need-to-know basis using the principle of least privilege. All access requests must be approved by the data owner and provisioned by IT Security. Access reviews are conducted quarterly."),
            ("Data Protection", "Confidential and Restricted data must be encrypted at rest and in transit. Portable storage devices must use approved encryption. Data loss prevention (DLP) tools monitor and prevent unauthorized data transfers."),
            ("Incident Response", "Security incidents must be reported to the IT Security team immediately. The Incident Response Plan defines procedures for containment, eradication, recovery, and post-incident review."),
            ("Compliance", "The information security program is designed to comply with applicable regulations including HIPAA, PCI-DSS, and state privacy laws. Annual security assessments and penetration testing are conducted."),
        ],
    },
    "information systems roles": {
        "purpose": "To define roles and responsibilities for information systems management and governance.",
        "scope": "This policy applies to all employees involved in the management, operation, and use of information systems.",
        "sections": [
            ("Chief Information Officer (CIO)", "The CIO is responsible for the overall strategy, direction, and management of the organization's information technology resources. The CIO reports to the executive leadership team."),
            ("IT Security Officer", "The IT Security Officer is responsible for implementing and maintaining the information security program, conducting risk assessments, and ensuring compliance with security policies."),
            ("System Administrators", "System administrators are responsible for the daily operation, maintenance, and security of assigned information systems. They must maintain current certifications and follow change management procedures."),
            ("End Users", "All employees are responsible for using information systems in accordance with established policies, protecting their access credentials, and reporting security concerns."),
        ],
    },
    "artificial intelligence|ai.*llm|generative": {
        "purpose": "To establish guidelines for the responsible use of generative artificial intelligence (AI) and large language models (LLMs) within the organization.",
        "scope": "This policy applies to all employees, contractors, and authorized users who use or interact with generative AI tools in the course of their work.",
        "sections": [
            ("Approved AI Tools", "Only AI tools approved by the IT Security team and listed in the Approved Software Registry may be used for business purposes. Employees must not use personal AI accounts for work-related tasks."),
            ("Data Protection", "Confidential, Restricted, or personally identifiable information (PII) must not be entered into external AI tools. Employees must review AI-generated content for accuracy before using it in business communications or decisions."),
            ("Content Review and Accountability", "AI-generated content must be reviewed and validated by the responsible employee before use. Employees remain accountable for the accuracy and appropriateness of any content they produce, whether AI-assisted or not."),
            ("Intellectual Property", "Employees should be aware of intellectual property implications when using AI tools. AI-generated content used in official documents must comply with the organization's intellectual property policies."),
            ("Training Requirements", "All employees using AI tools must complete the approved AI Literacy training course within 30 days of initial use and participate in annual refresher training."),
        ],
    },
    "mobile device": {
        "purpose": "To establish policies governing the use of mobile devices for accessing organizational data and systems.",
        "scope": "This policy applies to all organization-owned and personally-owned mobile devices used to access organizational resources.",
        "sections": [
            ("Device Registration", "All mobile devices used to access organizational email, applications, or data must be registered with IT and enrolled in the Mobile Device Management (MDM) system. Registration must occur before accessing any organizational resources."),
            ("Security Requirements", "Mobile devices must have: screen lock with PIN/biometric, encryption enabled, current operating system and security patches, approved antivirus software (if applicable), and remote wipe capability enabled."),
            ("Acceptable Use", "Mobile devices may be used to access organizational email, calendar, and approved applications. Users must not store Confidential or Restricted data on mobile devices without encryption. Lost or stolen devices must be reported to IT immediately."),
            ("BYOD (Bring Your Own Device)", "Employees using personal devices for work must acknowledge the organization's right to manage and remotely wipe organizational data on the device. Personal devices must meet the same security requirements as organization-issued devices."),
        ],
    },
    "configuration management": {
        "purpose": "To establish the configuration management process for IT systems and infrastructure.",
        "scope": "This policy applies to all IT systems, applications, and infrastructure components managed by the organization.",
        "sections": [
            ("Configuration Baselines", "All production systems must have documented configuration baselines. Baselines include hardware specifications, operating system configuration, installed software, network configuration, and security settings."),
            ("Change Control", "Changes to production configurations must follow the Change Management process. All changes require a change request, impact assessment, approval, and post-implementation review."),
            ("Configuration Monitoring", "Automated tools monitor production systems for configuration drift from approved baselines. Unauthorized configuration changes are flagged and investigated by IT Security."),
            ("Documentation", "Configuration documentation must be maintained in the Configuration Management Database (CMDB) and updated within 48 hours of any approved change."),
        ],
    },
    "cyber security committee": {
        "purpose": "To establish the charter and governance structure for the Cyber Security Committee.",
        "scope": "This charter defines the membership, responsibilities, and operating procedures of the Cyber Security Committee.",
        "sections": [
            ("Committee Mission", "The Cyber Security Committee provides strategic oversight of the organization's cybersecurity program, reviews security risks and incidents, and makes recommendations to executive leadership on security investments and policy."),
            ("Membership", "The Committee consists of the CIO (Chair), IT Security Officer, representatives from Legal, HR, Finance, and Operations, and external cybersecurity advisors as needed. Members serve two-year terms."),
            ("Meeting Schedule", "The Committee meets quarterly, with additional meetings called as needed for urgent security matters. Meeting minutes are distributed to all members within five business days."),
            ("Responsibilities", "The Committee reviews the annual cybersecurity risk assessment, approves the security awareness training program, reviews major security incidents, oversees compliance with regulatory requirements, and recommends the annual cybersecurity budget."),
        ],
    },
    "sop.*uniform|uniform.*issuance|uniform.*polo": {
        "purpose": "To provide standard operating procedures for the issuance, tracking, and return of employee uniforms.",
        "scope": "This SOP applies to all supervisors, uniform coordinators, and employees in positions requiring uniforms.",
        "sections": [
            ("Initial Uniform Issuance", "New employees receive uniform items as specified for their job classification within five (5) business days of hire. The uniform coordinator records all items issued in the Uniform Tracking System, including size, quantity, and date issued."),
            ("Replacement Procedures", "Employees may request replacement uniform items through their supervisor. Worn or damaged items must be returned when replacements are issued. Annual replacement allotments are determined by job classification."),
            ("Polo Shirt Specifications", "Approved polo shirts must be organization-branded with the official logo. Colors are department-specific as follows: Operations (navy blue), Administration (white), Field Services (grey). Polo shirts are worn on designated casual Fridays."),
            ("Return and Separation", "Employees must return all uniform items upon separation from employment. The cost of unreturned items may be deducted from the final paycheck in accordance with applicable state law."),
        ],
    },
    "manager toolkit": {
        "purpose": "To provide managers with a centralized resource for HR tools, templates, and reference materials.",
        "scope": "This document provides quick-reference links and resources for all people managers.",
        "sections": [
            ("Manager Resources", "The Manager Toolkit is available on the HR SharePoint site at the designated internal URL. It contains templates for performance reviews, corrective action, leave requests, and onboarding checklists."),
            ("Key HR Contacts", "For employee relations issues, contact the HR Generalist assigned to your department. For benefits questions, contact the Benefits Office. For payroll inquiries, contact Payroll Services. Contact information is maintained on the HR intranet."),
            ("Common Processes", "Performance reviews are conducted annually in Q4. Mid-year check-ins are required by June 30. PTO requests are managed through the Time-Off System. Hiring requests are submitted through the Applicant Tracking System."),
        ],
    },
    "blood.*borne.*pathogen.*intro": {
        "purpose": "To provide an introduction to the organization's Bloodborne Pathogens Exposure Control Program in compliance with OSHA Standard 29 CFR 1910.1030.",
        "scope": "This policy applies to all employees who have occupational exposure to blood or other potentially infectious materials (OPIM).",
        "sections": [
            ("Overview", "Bloodborne pathogens are pathogenic microorganisms present in human blood that can cause disease. The most significant pathogens include Hepatitis B Virus (HBV), Hepatitis C Virus (HCV), and Human Immunodeficiency Virus (HIV)."),
            ("Exposure Determination", "Job classifications with occupational exposure are identified in the Exposure Control Plan. Employees in these classifications must receive training within 10 days of assignment and annually thereafter."),
            ("Training Requirements", "All employees with occupational exposure must complete initial Bloodborne Pathogens training and annual refresher training. Training covers: epidemiology, modes of transmission, exposure control plan, PPE use, and post-exposure procedures."),
            ("Hepatitis B Vaccination", "The organization offers the Hepatitis B vaccination series at no cost to employees with occupational exposure. Employees may decline vaccination by signing a declination form, but may request vaccination at any time during employment."),
        ],
    },
    "blood.*borne.*pathogen.*compliance|blood.*borne.*pathogen.*method": {
        "purpose": "To define methods of compliance with the Bloodborne Pathogens Standard including engineering controls, work practice controls, and personal protective equipment.",
        "scope": "This policy applies to all employees with occupational exposure to bloodborne pathogens as identified in the Exposure Control Plan.",
        "sections": [
            ("Universal Precautions", "All employees must observe Universal Precautions, treating all human blood and OPIM as if known to be infectious. This is the foundation of the organization's Exposure Control Program."),
            ("Engineering Controls", "Engineering controls that reduce exposure include: sharps disposal containers, self-sheathing needles, needleless IV systems, biohazard labels, and hand-washing facilities. Engineering controls are inspected and maintained on a regular schedule."),
            ("Work Practice Controls", "Work practice controls include: hand washing immediately after glove removal, no eating or drinking in exposure areas, prohibiting mouth pipetting, minimizing splashing, and proper specimen handling procedures."),
            ("Personal Protective Equipment (PPE)", "Appropriate PPE is provided at no cost and includes gloves, gowns, face shields, masks, and eye protection. PPE must be used whenever there is potential for exposure to blood or OPIM."),
            ("Housekeeping and Decontamination", "Work areas are cleaned and decontaminated on a regular schedule and immediately after any spill of blood or OPIM using an appropriate disinfectant solution."),
        ],
    },
    "emergency notification": {
        "purpose": "To establish procedures for the organization's Emergency Notification System used to communicate critical information during emergencies.",
        "scope": "This policy applies to all employees and the administration of the Emergency Notification System.",
        "sections": [
            ("System Overview", "The Emergency Notification System (ENS) provides rapid multi-channel communication during emergencies via text message, email, phone call, and mobile app push notification."),
            ("Employee Enrollment", "All employees must register their current contact information in the ENS within 30 days of hire. Employees are responsible for keeping their contact information current."),
            ("Activation Criteria", "The ENS may be activated for: natural disasters, severe weather, facility closures, active threat situations, IT system outages, and other events that affect employee safety or operations."),
            ("Testing", "The ENS is tested quarterly. All employees will receive a test notification and must confirm receipt. Test results are reviewed by the Emergency Management Team."),
        ],
    },
    "computer replacement": {
        "purpose": "To establish the computer hardware replacement cycle and procedures for the organization.",
        "scope": "This policy applies to all organization-owned desktop computers, laptops, and related peripherals.",
        "sections": [
            ("Replacement Cycle", "Standard desktop computers and laptops are replaced on a four (4) year cycle. Specialized workstations may have extended or shortened cycles based on performance requirements and departmental needs."),
            ("Request Process", "Computer replacement requests are submitted through the IT Service Desk. Requests outside the standard cycle require justification and approval from the department head and IT Director."),
            ("Data Migration", "IT is responsible for migrating user data, profiles, and applications to replacement hardware. Users must ensure all critical files are saved to network drives or approved cloud storage prior to equipment swap."),
            ("Disposal", "Retired equipment is sanitized in accordance with NIST SP 800-88 guidelines and disposed of through the organization's approved electronics recycling program. Hard drives are physically destroyed."),
        ],
    },
    "technology master plan": {
        "purpose": "To define the organization's strategic technology roadmap and investment priorities for the planning period.",
        "scope": "This plan encompasses all information technology systems, infrastructure, and services across the organization.",
        "sections": [
            ("Strategic Vision", "The Integrated Technology Master Plan aligns IT investments with organizational strategic goals. Key focus areas include digital transformation, cybersecurity enhancement, data analytics capabilities, and workforce technology modernization."),
            ("Infrastructure Modernization", "Priority infrastructure projects include: cloud migration of core business applications, network infrastructure refresh, data center consolidation, and implementation of zero-trust security architecture."),
            ("Digital Services", "Planned digital service improvements include: employee self-service portal, mobile workforce enablement, automated workflow systems, and enhanced business intelligence and reporting platforms."),
            ("Budget and Timeline", "The five-year technology investment plan allocates resources across infrastructure, applications, security, and innovation. Annual budgets are reviewed and adjusted based on organizational priorities and technology evolution."),
            ("Governance", "The Technology Steering Committee provides oversight of master plan execution, reviews project portfolios quarterly, and ensures alignment with organizational strategic objectives."),
        ],
    },
    "information technology mission": {
        "purpose": "To articulate the mission, vision, and guiding principles of the Information Technology department.",
        "scope": "This statement applies to the IT department and defines its commitment to the organization.",
        "sections": [
            ("Mission Statement", "The Information Technology department delivers reliable, secure, and innovative technology solutions that enable the organization to achieve its mission and serve its stakeholders effectively."),
            ("Vision", "To be a trusted technology partner that drives operational excellence and enables data-driven decision-making across the organization through modern, secure, and user-friendly systems."),
            ("Core Values", "Innovation: We embrace emerging technologies to improve service delivery. Security: We protect organizational assets and data. Service: We provide responsive and professional IT support. Collaboration: We partner with departments to understand and address their technology needs."),
        ],
    },
}


def _match_policy(filename: str) -> dict | None:
    """Return the best-matching content template for a filename."""
    name_lower = filename.lower()
    for pattern, content in POLICY_CONTENT.items():
        if re.search(pattern, name_lower):
            return content
    return None


def _parse_title(filename: str) -> tuple[str, str]:
    """Extract policy number and title from filename like '51350 - Types of Leave...'."""
    stem = Path(filename).stem
    # Remove trailing (12345_N) version info
    stem = re.sub(r"\s*\(\d+_\d+\)\s*$", "", stem)
    m = re.match(r"^(\d+)\s*[-–]\s*(.+)$", stem)
    if m:
        return m.group(1), m.group(2).strip()
    return "", stem.strip()


def generate_docx(filepath: Path, dry_run: bool = False) -> bool:
    """Generate a proper .docx file with synthetic HR policy content."""
    filename = filepath.name
    policy_number, title = _parse_title(filename)
    content = _match_policy(filename)

    if not content:
        logger.warning("No template match for '%s' — generating generic content", filename)
        content = {
            "purpose": f"To establish policies and procedures related to {title}.",
            "scope": "This policy applies to all employees of the organization.",
            "sections": [
                ("Policy Statement", f"The organization maintains standards and procedures for {title} in accordance with applicable laws, regulations, and industry best practices."),
                ("Responsibilities", f"Department heads and supervisors are responsible for ensuring compliance with {title} requirements within their areas of responsibility."),
                ("Procedures", f"Detailed procedures for {title} are maintained by the responsible department and reviewed annually."),
                ("Compliance", "Violations of this policy may result in disciplinary action up to and including termination of employment."),
            ],
        }

    if dry_run:
        logger.info("[DRY RUN] Would regenerate: %s", filepath)
        return True

    doc = Document()

    # Title
    heading = f"Policy {policy_number} — {title}" if policy_number else title
    doc.add_heading(heading, level=0)

    # Purpose
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(content["purpose"])

    # Scope
    doc.add_heading("Scope", level=1)
    doc.add_paragraph(content["scope"])

    # Sections
    for section_title, section_body in content["sections"]:
        doc.add_heading(section_title, level=2)
        doc.add_paragraph(section_body)

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"This document is the property of the organization. "
        f"Policy {policy_number}. " if policy_number else ""
    ).italic = True

    doc.save(str(filepath))
    logger.info("Regenerated: %s", filepath)
    return True


def main():
    parser = argparse.ArgumentParser(description="Regenerate CDFV2-encrypted .docx with synthetic content")
    parser.add_argument("--dry-run", action="store_true", help="List files that would be regenerated")
    args = parser.parse_args()

    if not DATA_DIR.exists():
        logger.error("Data directory not found: %s", DATA_DIR)
        sys.exit(1)

    encrypted = []
    for docx_path in sorted(DATA_DIR.rglob("*.docx")):
        result = subprocess.run(["file", "-b", str(docx_path)], capture_output=True, text=True)
        if "CDFV2 Encrypted" in result.stdout:
            encrypted.append(docx_path)

    # Also check .xlsx
    for xlsx_path in sorted(DATA_DIR.rglob("*.xlsx")):
        result = subprocess.run(["file", "-b", str(xlsx_path)], capture_output=True, text=True)
        if "CDFV2 Encrypted" in result.stdout:
            encrypted.append(xlsx_path)

    logger.info("Found %d CDFV2 Encrypted file(s)", len(encrypted))

    regenerated = 0
    skipped = 0
    for filepath in encrypted:
        if filepath.suffix == ".xlsx":
            logger.info("Skipping .xlsx (not a policy document): %s", filepath.name)
            skipped += 1
            continue
        if generate_docx(filepath, dry_run=args.dry_run):
            regenerated += 1
        else:
            skipped += 1

    logger.info("Done — %d regenerated, %d skipped", regenerated, skipped)


if __name__ == "__main__":
    main()
